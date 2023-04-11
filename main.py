import tkinter as tk
from tkinter import filedialog
from PIL import ImageTk, Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import Table
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from docx.oxml import parse_xml


def browse_file():
    # Open a file dialog to select an image file
    file_path = filedialog.askopenfilename()
    if file_path:
        # Open the selected image file
        img = Image.open(file_path)
        # Resize the image to fit in the Word document
        img.thumbnail((1.5, 1.5))
        # Convert the image to RGB mode
        img = img.convert("RGB")
        # Save the image to a temporary file
        img_path = "temp_logo.jpg"
        img.save(img_path)
        # Update the logo label with the selected image
        logo_label.config(text=f"Logo: {file_path}")
        logo_label.image = ImageTk.PhotoImage(Image.open(img_path))
        

def table_title_color(tbl_name, row, col):
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="6baede"/>'.format(nsdecls('w')))
    tbl_name.rows[row].cells[col]._tc.get_or_add_tcPr().append(shading_elm_1)
    tbl_name.cell(row, col).paragraphs[0].runs[0].bold = True  # Set text to bold

def create_doc_info_table(doc, title):
    # Create a table for document information with all borders
    
    doc.add_heading('Document Information', level=1)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'  # Set table style to "Table Grid" which includes all borders
    

    # Populate the table with data
    info_table.cell(0, 0).text = 'Category'
    info_table.cell(0, 1).text = 'Information'
    info_table.cell(1, 0).text = 'Document'
    info_table.cell(1, 1).text = title
    info_table.cell(2, 0).text = 'Document Version'
    info_table.cell(2, 1).text = '1.0'
    info_table.cell(3, 0).text = 'Classification Level'
    info_table.cell(3, 1).text = 'Confidential'
    info_table.cell(4, 0).text = 'Author(s)'
    info_table.cell(4, 1).text = ''
    
    table_title_color(info_table, 0, 0)
    table_title_color(info_table, 0, 1)
    #info_table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    #info_table.cell(0, 0).paragraphs[0].runs[0].bold = True  # Set text to bold
    

    # Set alignment of the table to left
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Add a page break

    # Add document revision history as heading
    doc.add_heading('Document Revision History', level=1)
    # Create a table for document revision history
    rev_table = doc.add_table(rows=4, cols=4)
    rev_table.style = 'Table Grid'  # Set table style to "Table Grid" which includes all borders
    # Populate the table with data
    rev_table.cell(0, 0).text = 'Author'
    rev_table.cell(0, 1).text = 'Date'
    rev_table.cell(0, 2).text = 'Version'
    rev_table.cell(0, 3).text = 'Description'
    # Add 3 blank rows
    for i in range(1, 4):
        for j in range(4):
            rev_table.cell(i, j).text = ''

    # Set alignment of the table to left
    rev_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set row height for revision history table
    for row in rev_table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        row.height = Pt(12)

def create_word_file():
    # Retrieve input from GUI input fields
    title = 'Application Security Assessment - '+ title_entry.get()
    
    
    content = content_entry.get(1.0, tk.END)

    # Create a new Word document
    doc = Document()

# Set the title as the document's title
    doc_tilte = doc.add_heading('\n\n\n\n'+title, 0)
    doc_tilte.style.font.size = Pt(48)
    # Add a title page with the logo
    #if logo_label.image:
    #    doc.add_picture("temp_logo.jpg", width=Inches(1.5))
    #doc.add_heading(title, 0)

    # Add the content to the document
    doc.add_paragraph(content)
    

    # Delete the temporary image file
    # if logo_label.image:
        # logo_label.image = None
        # logo_label.config(text="Logo: None")
        # img_path = "temp_logo.jpg"
        # import os
        # if os.path.exists(img_path):
            # os.remove(img_path)
	# Add document information as heading
    doc.add_page_break()
    create_doc_info_table(doc, title)
    
    # Save the document with the title as the filename
    doc.save(f"{title}.docx")
    status_label.config(text=f"Word file '{title}.docx' created successfully!")


# Create a tkinter window
window = tk.Tk()
window.title("Create Word File")

# Create input labels
title_label = tk.Label(window, text="Title:")
title_label.grid(row=0, column=0, sticky=tk.W)
content_label = tk.Label(window, text="Content:")
content_label.grid(row=1, column=0, sticky=tk.W)

# Create input fields
title_entry = tk.Entry(window)
title_entry.grid(row=0, column=1, columnspan=2)
content_entry = tk.Text(window, height=5, width=30)
content_entry.grid(row=1, column=1, columnspan=2)

# Create logo label
logo_label = tk.Label(window, text="Logo: None")
logo_label.grid(row=2, column=0, sticky=tk.W)

# Create browse button for logo
browse_button = tk.Button(window, text="Browse Logo", command=browse_file)
browse_button.grid(row=2, column=1)

# Create create button
create_button = tk.Button(window, text="Create Word File", command=create_word_file)
create_button.grid(row=2, column=2)

# Create status label
status_label = tk.Label(window, text="")
status_label.grid(row=3, column=0, columnspan=3)

# Run the tkinter event loop
window.mainloop()
